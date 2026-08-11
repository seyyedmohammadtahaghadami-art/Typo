from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from PIL import Image

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    OxmlElement = None
    qn = None

FA_CHARS = set("اآبپتثجچحخدذرزژسشصضطظعغفقکگلمنوهی")
RTL_RE = re.compile(r"[\u0600-\u06FF]")

def has_persian(text: str) -> bool:
    return any(c in FA_CHARS for c in text)

def clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()

def add_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

def add_text(doc: Document, text: str):
    text = clean_text(text)
    if not text:
        return
    p = doc.add_paragraph()
    add_rtl(p) if has_persian(text) else None
    run = p.add_run(text)
    run.font.size = Pt(11)

def ocr_image(path: Path) -> str:
    if pytesseract is None:
        raise RuntimeError("Tesseract/PyTesseract روی سرور نصب نیست.")
    # Persian + English language pack. If Persian isn't installed, give a clear error.
    try:
        return pytesseract.image_to_string(Image.open(path), lang="fas+eng", config="--psm 6")
    except Exception as exc:
        raise RuntimeError(
            "OCR فارسی اجرا نشد. مطمئن شوید Tesseract و بسته زبان fas نصب شده‌اند."
        ) from exc

def convert_pdf(path: Path, doc: Document, ocr_mode: str):
    if fitz is None:
        raise RuntimeError("PyMuPDF نصب نیست.")
    pdf = fitz.open(path)
    for page in pdf:
        text = page.get_text("text") or ""
        if ocr_mode == "force" or (ocr_mode == "auto" and len(text.strip()) < 20):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            if pytesseract is None:
                raise RuntimeError("Tesseract برای OCR در دسترس نیست.")
            try:
                text = pytesseract.image_to_string(img, lang="fas+eng", config="--psm 6")
            except Exception as exc:
                raise RuntimeError("OCR فارسی/انگلیسی اجرا نشد.") from exc
        for line in text.splitlines():
            if clean_text(line):
                add_text(doc, line)

def convert_image(path: Path, doc: Document):
    text = ocr_image(path)
    for line in text.splitlines():
        if clean_text(line):
            add_text(doc, line)

def convert_to_docx(source: Path, ocr_mode="auto", table_mode="balanced", output_dir=None) -> Path:
    output_dir = Path(output_dir or source.parent)
    output = output_dir / "TYPO-Real-Converted.docx"
    doc = Document()

    if source.suffix.lower() == ".pdf":
        convert_pdf(source, doc, ocr_mode)
    else:
        convert_image(source, doc)

    if not doc.paragraphs:
        add_text(doc, "متنی از فایل قابل استخراج نبود.")

    doc.save(output)
    return output
