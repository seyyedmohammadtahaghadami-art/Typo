import io
import re
import unicodedata
from typing import Any

import fitz
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


def norm(s: str) -> str:
    s = unicodedata.normalize('NFC', s or '')
    s = (s.replace('ي', 'ی').replace('ى', 'ی').replace('ك', 'ک')
           .replace('\u200f', '').replace('\u200e', ''))
    return re.sub(r'\s+', ' ', s).strip()


def is_rtl(s: str) -> bool:
    return bool(re.search(r'[\u0600-\u06ff]', s or ''))


def set_run_font(run, font_name: str | None):
    if not font_name:
        return
    run.font.name = font_name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rfonts.set(qn(f'w:{attr}'), font_name)


def set_paragraph_rtl(p, rtl: bool):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if rtl and bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    elif not rtl and bidi is not None:
        pPr.remove(bidi)


def lines_from_words(words: list[dict[str, Any]], tol=7):
    groups = []
    for w in sorted(words, key=lambda z: (z['y'], z['x'])):
        g = next((g for g in groups if abs(g['y'] - w['y']) <= max(tol, w['h'] * 0.55)), None)
        if not g:
            g = {'y': w['y'], 'h': w['h'], 'words': []}
            groups.append(g)
        g['words'].append(w)
        g['h'] = max(g['h'], w['h'])
    out = []
    for g in sorted(groups, key=lambda z: z['y']):
        rr = sum(is_rtl(x['text']) for x in g['words']) > len(g['words']) / 2
        g['words'].sort(key=lambda x: x['x'], reverse=rr)
        out.append({'text': ' '.join(x['text'] for x in g['words']), 'rtl': rr, 'words': g['words'], 'y': g['y'], 'h': g['h']})
    return out


def detect_table(lines, mode='balanced'):
    if mode == 'off':
        return None
    # Conservative heuristic: repeated x anchors across multiple lines + enough rows.
    anchors = []
    for line in lines:
        for w in line['words']:
            anchors.append((round(w['x'] / 10) * 10, line['y']))
    counts = {}
    for x, _ in anchors:
        counts[x] = counts.get(x, 0) + 1
    min_count = 3 if mode == 'strict' else 2
    cols = sorted(x for x, c in counts.items() if c >= min_count)
    if len(cols) < 2:
        return None
    rows = []
    for line in lines:
        row = []
        for i, x in enumerate(cols):
            left = -1e9 if i == 0 else (cols[i - 1] + x) / 2
            right = 1e9 if i == len(cols) - 1 else (x + cols[i + 1]) / 2
            row.append(' '.join(w['text'] for w in line['words'] if left <= w['x'] < right).strip())
        if sum(bool(v) for v in row) >= 2:
            rows.append(row)
    return rows if len(rows) >= (3 if mode == 'strict' else 2) else None


def _quality(words):
    if not words:
        return 0.0
    chars = ''.join(w['text'] for w in words)
    printable = sum(c.isprintable() for c in chars)
    useful = sum(bool(re.search(r'[\w\u0600-\u06ff]', w['text'])) for w in words)
    return min(1.0, (printable / max(1, len(chars))) * 0.55 + (useful / len(words)) * 0.45)


def extract_page(page, ocr_mode, language):
    words = []
    data = page.get_text('dict', flags=fitz.TEXTFLAGS_TEXT)
    for block in data.get('blocks', []):
        for line in block.get('lines', []):
            for span in line.get('spans', []):
                text = norm(span.get('text', ''))
                if not text:
                    continue
                x0, y0, x1, y1 = span['bbox']
                words.append({
                    'text': text, 'x': x0, 'y': y0, 'w': x1 - x0,
                    'h': max(1, y1 - y0), 'font': span.get('font', ''),
                    'size': span.get('size', 11), 'flags': span.get('flags', 0), 'confidence': 100,
                })

    use_ocr = ocr_mode == 'always' or (ocr_mode == 'auto' and (_quality(words) < 0.72 or len(words) < 4))
    if use_ocr:
        pix = page.get_pixmap(matrix=fitz.Matrix(2.4, 2.4), alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes('png'))).convert('RGB')
        try:
            raw = pytesseract.image_to_data(img, lang=language, output_type=pytesseract.Output.DICT, config='--oem 1 --psm 3')
        except pytesseract.TesseractError:
            raw = pytesseract.image_to_data(img, lang='eng', output_type=pytesseract.Output.DICT, config='--oem 1 --psm 3')
        words = []
        for i, text in enumerate(raw.get('text', [])):
            text = norm(text)
            conf = float(raw.get('conf', ['0'])[i]) if raw.get('conf') else 0
            if not text or conf < 5:
                continue
            words.append({
                'text': text,
                'x': raw['left'][i] / 2.4,
                'y': raw['top'][i] / 2.4,
                'w': raw['width'][i] / 2.4,
                'h': max(1, raw['height'][i] / 2.4),
                'font': '', 'size': max(8, raw['height'][i] / 2.4), 'flags': 0,
                'confidence': conf,
            })
    return lines_from_words(words), use_ocr


def add_image_to_doc(doc, page, rect, image_bytes, page_width):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        width = min(max(rect.width / 72, 0.6), max(0.8, page_width - 1.0))
        p.add_run().add_picture(io.BytesIO(image_bytes), width=Inches(width))
    except Exception:
        pass


def convert_document(data, filename, ocr_mode='auto', language='fas+eng', table_mode='balanced'):
    if data[:4] == b'%PDF':
        pdf = fitz.open(stream=data, filetype='pdf')
        if pdf.page_count == 0:
            raise ValueError('PDF has no pages.')
    else:
        try:
            Image.open(io.BytesIO(data)).verify()
        except Exception as exc:
            raise ValueError('Invalid image file.') from exc
        img = Image.open(io.BytesIO(data)).convert('RGB')
        pdf = fitz.open()
        page = pdf.new_page(width=img.width, height=img.height)
        page.insert_image(page.rect, stream=data)

    doc = Document()
    # A4-like margins; Word remains editable while preserving readable spacing.
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    for pi, page in enumerate(pdf):
        lines, used_ocr = extract_page(page, ocr_mode, language)
        table = detect_table(lines, table_mode)
        if table:
            t = doc.add_table(rows=len(table), cols=max(map(len, table)))
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r, row in enumerate(table):
                for c, value in enumerate(row):
                    cell = t.cell(r, c)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p = cell.paragraphs[0]
                    set_paragraph_rtl(p, is_rtl(value))
                    p.add_run(value)

        for line in lines:
            p = doc.add_paragraph()
            set_paragraph_rtl(p, line['rtl'])
            p.paragraph_format.space_after = Pt(2)
            for w in line['words']:
                run = p.add_run(w['text'] + ' ')
                run.font.size = Pt(max(6, min(72, float(w.get('size', 11)))))
                set_run_font(run, w.get('font'))
                flags = int(w.get('flags', 0))
                run.bold = bool(flags & 16)
                run.italic = bool(flags & 2)

        # Preserve embedded raster images when the PDF exposes them.
        seen = set()
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in seen:
                continue
            seen.add(xref)
            try:
                extracted = pdf.extract_image(xref)
                if extracted.get('image'):
                    add_image_to_doc(doc, page, fitz.Rect(0, 0, page.rect.width, page.rect.height), extracted['image'], page.rect.width / 72)
            except Exception:
                continue

        if pi < len(pdf) - 1:
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    pdf.close()
    return out.getvalue()
